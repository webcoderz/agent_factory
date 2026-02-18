from __future__ import annotations
import io
from agent_ext.export.models import ExportRequest

class DocxExporter:
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def filename(self, *, req: ExportRequest) -> str:
        return "report.docx"

    def render_bytes(self, *, req: ExportRequest, outcome: dict) -> bytes:
        try:
            from docx import Document
        except Exception as e:
            raise ImportError("python-docx required: pip install python-docx") from e

        doc = Document()
        doc.add_heading(req.title, level=0)

        answer = str(outcome.get("answer", ""))
        if answer:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(answer)

        if req.include_claims:
            claims = outcome.get("claims") or []
            if claims:
                doc.add_heading("Claims", level=1)
                for c in claims:
                    txt = c.get("text") if isinstance(c, dict) else str(c)
                    doc.add_paragraph(str(txt), style="List Bullet")

        if req.include_limitations:
            limitations = outcome.get("limitations") or []
            if limitations:
                doc.add_heading("Limitations", level=1)
                for l in limitations:
                    doc.add_paragraph(str(l), style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
